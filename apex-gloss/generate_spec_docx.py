"""Generate 'XLETH — APEX & GLOSS Design Spec.docx' per the xleth-mentor style/structure contracts."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

STEEL = RGBColor(0x4F, 0x81, 0xBD)
BLACK = RGBColor(0, 0, 0)

OUT = r"C:\Users\Krasen\Documents\kimi\workspace\apex-gloss\XLETH — APEX & GLOSS Design Spec.docx"

doc = Document()

for s in doc.sections:
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1)

normal = doc.styles["Normal"]
normal.font.name = "DejaVu Serif"
normal.font.size = Pt(11)
normal.font.color.rgb = BLACK
normal.element.get_or_add_rPr()
rfonts = normal.element.rPr.get_or_add_rFonts()
rfonts.set(qn("w:eastAsia"), "DejaVu Serif")
normal.paragraph_format.space_after = Pt(6)

def style_heading(name, size, color=STEEL, bold=True, font="Carlito"):
    st = doc.styles[name]
    st.font.name = font
    st.font.size = Pt(size)
    st.font.bold = bold
    st.font.color.rgb = color
    rf = st.element.get_or_add_rPr().get_or_add_rFonts()
    rf.set(qn("w:ascii"), font)
    rf.set(qn("w:hAnsi"), font)
    rf.set(qn("w:eastAsia"), font)

style_heading("Heading 1", 16)
style_heading("Heading 2", 14)
style_heading("Heading 3", 12, color=BLACK)
style_heading("Heading 4", 11, color=BLACK)

def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def para(lead=None, text=""):
    p = doc.add_paragraph()
    if lead:
        r = p.add_run(lead)
        r.bold = True
    if text:
        p.add_run(text)
    return p

def bullet(lead=None, text=""):
    p = doc.add_paragraph(style="List Bullet")
    if lead:
        r = p.add_run(lead)
        r.bold = True
    if text:
        p.add_run(text)
    return p

def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    pPr.append(pbdr)

def code_block(text):
    for line in text.strip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run(line if line else " ")
        r.font.name = "Noto Sans Mono"
        r.font.size = Pt(10)
        rf = r._element.get_or_add_rPr().get_or_add_rFonts()
        rf.set(qn("w:ascii"), "Noto Sans Mono")
        rf.set(qn("w:hAnsi"), "Noto Sans Mono")

def set_cell(cell, text, bold=False, mono=False):
    cell.paragraphs[0].text = ""
    r = cell.paragraphs[0].add_run(text)
    r.bold = bold
    r.font.size = Pt(11)
    if mono:
        r.font.name = "Noto Sans Mono"
        r.font.size = Pt(10)

def add_table(headers, rows):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tbl = t._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    for edge in ("left", "right", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        borders.append(el)
    tblPr.append(borders)
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, bold=True)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            mono = ci == 0
            set_cell(t.rows[ri].cells[ci], val, mono=mono)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t

# ============================ CONTENT ============================

h1("XLETH — APEX & GLOSS Design Spec")

para("What this is: ",
     "the complete design specification for APEX, XLETH's native multiband maximizer stock effect "
     "(functionally modeled on Image-Line Maximus), and GLOSS, its one-knob companion skin "
     "(functionally modeled on Soundgoodizer). This document is the single source of truth for the "
     "four Claude Code implementation prompts that accompany it.")
para("Trigger: ",
     "any engineering work on APEX or GLOSS. Load this before modifying DSP, bridge, or UI code for "
     "these effects.")

hr()

h2("1. Overview")

bullet("APEX: ",
       "a native stock effect in XLETH's Phase 3 effects system — three user-definable bands "
       "(LOW / MID / HIGH) plus a MASTER wideband stage, each a full dynamics processor with a "
       "node-based curve editor, saturation, stereo separation, and gain staging. Not a VST3; it "
       "lives inside the C++ engine as an AudioGraph effect node with a React editor window.")
bullet("GLOSS: ",
       "a second UI descriptor over the identical DSP engine: one giant AMOUNT knob (bound to BAND "
       "MIX) plus four lettered preset buttons (A–D). One DSP class, two skins — every APEX fix "
       "propagates to GLOSS for free.")
bullet("Design source: ",
       "Maximus multiband maximizer (Image-Line) for APEX's feature set; Soundgoodizer for GLOSS's "
       "interaction model. Feature-parity where it matters, simplified where Maximus is over-built "
       "(no REL 2, no envelope tension knobs, no linear-phase crossover mode in v1).")

hr()

h2("2. Naming & IP boundaries")

bullet("APEX ", "= the Maximus analog. Peak of loudness; reads instantly as a maximizer to anyone "
       "coming from FL Studio. Distinct from the Maximus trademark.")
bullet("GLOSS ", "= the Soundgoodizer analog. \"It's shiny, it's good\" — gloss is shine. Header "
       "reads \"powered by APEX\", preserving the parent/child family link.")
bullet("Renamed parameters: ", "LMH MIX → BAND MIX (in GLOSS: AMOUNT). LMH DEL → LOOKAHEAD. All "
       "other labels (LOW/MID/HIGH/MASTER, SOLO, PRE/POST GAIN, ATT/REL/SUSTAIN, PEAK/RMS, THRESH, "
       "CEIL, STEREO SEP, LOW CUT, 12/24 dB, ON/COMP OFF/MUTED/OFF) are generic industry vocabulary "
       "and stay as-is.")
bullet("Hard IP rules: ",
       "DSP concepts (multiband compression, saturation, lookahead, M/S width) are not protectable "
       "and are safe to implement. Do NOT copy Image-Line artwork, manual text, or FL preset files. "
       "GLOSS's A–D preset values must be authored in-house by ear — the letters are generic, the "
       "values are Image-Line's. Do not attempt to extract them.")

hr()

h2("3. Signal flow")

code_block("""
Input
  └─ LOW CUT (HPF, 0–100 Hz, 24 dB/oct)
       └─ Crossover: LR2 (12 dB/oct) or LR4 (24 dB/oct), per split
          ├─ LOW band   ┐
          ├─ MID band   ├─ per band: PRE GAIN → LOOKAHEAD delay
          └─ HIGH band  ┘        → envelope detect (PEAK/RMS, ATT/REL/SUSTAIN)
                                 → curve-LUT gain → SAT (CEIL + bipolar A/B, 2× oversampled)
                                 → POST GAIN → STEREO SEP (M/S width)
                                 → band state (ON / COMP OFF / MUTED / OFF)
       └─ Σ bands  ⇄  dry input (delayed to match LOOKAHEAD) via BAND MIX
            └─ MASTER band (same chain, wideband, no lookahead)
                 └─ Output
""")

bullet("Band state semantics (exact): ",
       "ON = full chain. COMP OFF = dynamics bypassed; saturation, gains, and separation still "
       "active; band passes. MUTED = band silenced. OFF = band bypassed dry and its DSP skipped "
       "(CPU saved).")
bullet("Dry-path alignment: ",
       "the dry side of BAND MIX must be delayed by exactly the LOOKAHEAD amount so the parallel "
       "blend is time-aligned. LR crossovers sum magnitude-flat, which is what makes phase-free "
       "parallel blending possible at all.")
bullet("MASTER stage: ",
       "processes the post-mix signal. No lookahead on MASTER in v1 — LOOKAHEAD covers L/M/H only, "
       "per spec.")

hr()

h2("4. Parameter specification")

h3("4.1 Global & crossover")
add_table(
    ["Parameter", "Range", "Default", "Notes"],
    [
        ["LOW CUT", "0–100 Hz", "0 (off)", "HPF before the crossover, 24 dB/oct"],
        ["LOW split", "40 Hz–1 kHz", "200 Hz", "LOW/MID crossover frequency"],
        ["HIGH split", "1 kHz–18 kHz", "2 kHz", "MID/HIGH crossover frequency"],
        ["LOW slope", "12 / 24 dB/oct", "24 dB", "LR2 / LR4, per-split selectable"],
        ["HIGH slope", "12 / 24 dB/oct", "24 dB", "LR2 / LR4, per-split selectable"],
        ["LOOKAHEAD", "0–20 ms", "0 ms", "Shared delay for L/M/H; reported as effect latency"],
        ["BAND MIX", "0–100 %", "100 %", "Parallel blend: dry input ↔ summed L/M/H output"],
    ],
)

h3("4.2 Per-band (identical for LOW / MID / HIGH / MASTER)")
add_table(
    ["Parameter", "Range", "Default", "Notes"],
    [
        ["STATE", "ON / COMP OFF / MUTED / OFF", "ON", "4-position switch, semantics in Section 3"],
        ["SOLO", "toggle", "off", "L/M/H only; mutually exclusive across bands"],
        ["PRE GAIN", "-24…+24 dB", "0 dB", "Before dynamics"],
        ["POST GAIN", "-24…+24 dB", "0 dB", "After saturation / separation"],
        ["CURVE", "node editor", "unity (flat)", "Up to 32 nodes, per-segment tension"],
        ["ATT", "0.1–100 ms", "5 ms", "Envelope attack"],
        ["REL", "5–500 ms", "100 ms", "Envelope release"],
        ["SUSTAIN", "0–500 ms", "0 ms", "Envelope hold before release starts"],
        ["Detection", "PEAK / RMS", "PEAK", "Envelope detector mode"],
        ["SAT THRESH", "-100…+100 %", "0 %", "<0 = mode A, 0 = dry, >0 = mode B"],
        ["SAT CEIL", "-60…0 dB", "0 dB", "Saturation onset level"],
        ["STEREO SEP", "-100…+100 %", "0 %", "M/S width; <0 narrows toward mono"],
    ],
)

h3("4.3 Curve editor")
bullet("Node model: ",
       "each node is (IN dB, OUT dB); both axes -24…+12 dB. Up to 32 nodes per band. Nodes are "
       "sorted by IN; first and last nodes pin the curve ends.")
bullet("Per-segment tension: ",
       "each segment carries a tension value -1…+1 (0 = linear). Tension is what enables soft "
       "knees versus brickwall shapes — compression, limiting, expansion, and gating all come out "
       "of this one editor.")
bullet("Editing operations: ",
       "drag to move a node, double-click to add, right-click to delete, drag/wheel on a segment "
       "handle to set tension.")
bullet("LUT compilation: ",
       "on any node or tension edit, the engine rebuilds a 1024-entry dB→dB gain LUT off the "
       "audio thread and hands it to the audio thread by atomic pointer swap. The audio callback "
       "only ever does LUT lookups.")

hr()

h2("5. DSP implementation requirements")

bullet("Crossover: ",
       "Linkwitz-Riley 2nd order (12 dB/oct) and 4th order (24 dB/oct, two cascaded 2nd-order "
       "Butterworth), selectable per split. Magnitude-flat reconstruction is a hard requirement — "
       "verified by test, not assumed. No linear-phase FIR mode in v1 (latency cost, no benefit in "
       "a Sparta workflow).")
bullet("Envelope detector: ",
       "branching attack/release smoother with hold; PEAK or RMS detection selectable per band. "
       "SUSTAIN holds the gain reduction at its current value before release begins.")
bullet("Saturation: ",
       "per band including MASTER. CEIL sets the onset level: signal exceeding the ceiling gets "
       "shaped, scaled by the THRESH amount. Mode A = smooth soft-clip (tanh-family); mode B = "
       "harder clip with richer high-order harmonics. The waveshaper — and only the waveshaper — "
       "runs under 2× oversampling (half-band polyphase up/down) to prevent aliasing on "
       "percussion material.")
bullet("Stereo separation: ",
       "per-band M/S encode → side gain scale → decode. -100 % = mono, 0 = unchanged, +100 % = "
       "double side energy. Applied to MASTER as well.")
bullet("Lookahead: ",
       "shared pre-allocated ring-buffer delay for L/M/H, 0–20 ms, smoothed on change (no zipper "
       "noise, no reallocation).")
bullet("Audio-thread rules (hard constraints): ",
       "no allocations, no locks, no logging on the audio callback path. All parameter changes "
       "arrive through the existing parameter queue; the audio thread consumes pre-built state "
       "and atomically-swapped LUT pointers only.")

hr()

h2("6. Engine integration")

bullet("Effect node: ",
       "APEX registers as a standard stock effect in the Phase 3 effects system — same lifecycle, "
       "parameter API, and state serialization path as existing stock effects. UndoManager and "
       "project save/load must work unchanged.")
bullet("Latency (the one engine-level risk): ",
       "LOOKAHEAD makes APEX a latency-introducing effect. The implementation must audit what "
       "latency reporting / delay compensation exists in the engine today. If none exists, APEX "
       "reports its latency through a clean new mechanism consistent with the architecture, the "
       "UI displays it, and the exact monitoring-delay behavior is documented. Do not hand-wave "
       "this.")
bullet("Metering: ",
       "per-band gain reduction (dB), per-band output level, and input spectrum (FFT 2048) pushed "
       "to the UI at ~30 Hz as one batched typed-array payload per tick — never per-scalar RPC "
       "per frame. Reuse the existing metering/telemetry path if one exists.")
bullet("Bridge discipline: ",
       "the four-layer bridge silently swallows undefined calls via optional chaining — every new "
       "bridge call must be verified end-to-end with console smoke tests on both sides. "
       "child_process.fork must use serialization: 'advanced' so metering ArrayBuffers survive "
       "intact. build.bat bridge-clean is mandatory after any C++ change before concluding "
       "anything works or does not.")

hr()

h2("7. UI specification")

h3("7.1 APEX editor window")
bullet("Layout: ",
       "curve editor (left) + analysis display (right, spectrum with band regions tinted per "
       "band) on top; band selector tabs (LOW / MID / HIGH / MASTER) with per-band 4-state switch "
       "and SOLO below; knob rows for PRE/POST GAIN, ATT/REL/SUSTAIN + PEAK/RMS toggle, SAT "
       "THRESH/CEIL, STEREO SEP; right panel for LOOKAHEAD, BAND MIX, LOW/HIGH split + slope "
       "switches, LOW CUT.")
bullet("Canvas discipline: ",
       "curve editor and analysis display are canvas-drawn — never DOM overlays over canvas "
       "(scroll/zoom drift, hit-test corruption).")
bullet("Drag discipline: ",
       "local preview during drag, single commit on mouseup through UndoManager. Never IPC "
       "during drag.")
bullet("Design system: ",
       "CSS custom-property tokens throughout, no hardcoded hex; tokenValue() never at module "
       "scope. \"Color is earned, not default\" — accent color only on meaningful state (active "
       "band, gain-reduction activity, solo/mute). Lucide icons only. Flat, zero-border-radius, "
       "teal accent hierarchy per the established theming spec.")
bullet("Script safety: ",
       "never declare reserved identifiers (top, self, parent, name, length, status, location, "
       "history, event, screen, navigator, opener, origin, closed) as variables in script scope.")

h3("7.2 GLOSS window")
bullet("Controls: ",
       "one giant AMOUNT knob bound to BAND MIX (0–100 %), four lettered buttons (A / B / C / D) "
       "loading APEX parameter snapshots, header \"powered by APEX\".")
bullet("Presets: ",
       "A–D ship as four factory snapshots authored in-house (suggested starting points: gentle "
       "glue, punchy, bright lift, wall). Each is a complete APEX parameter state. AMOUNT stays "
       "independent of preset selection.")
bullet("Architecture: ",
       "GLOSS is a second UI descriptor over the same engine class — not a separate effect, not "
       "duplicated DSP.")

hr()

h2("8. Implementation order")

para("",
     "Four Claude Code prompts, executed strictly in order. Each is one logical unit of work with "
     "its own build verification and git commit. Do not batch them.")
bullet("Prompt 1 — DSP core (Opus, High): ",
       "audit the existing stock-FX framework, then implement the APEX engine class: crossover, "
       "per-band chain, curve LUT, envelope, oversampled saturation, separation, band states, "
       "lookahead + latency reporting, BAND MIX with aligned dry path. Verified by engine-side "
       "console smoke tests.")
bullet("Prompt 2 — Bridge & metering (Opus, High): ",
       "wire all parameters, curve data, and state through the Node-API bridge end-to-end; "
       "batched metering channel; latency surfaced to the UI.")
bullet("Prompt 3 — APEX editor UI (Sonnet, High): ",
       "canvas curve editor with nodes + segment tension, canvas analysis display, full control "
       "set, Zustand store, drag/undo discipline.")
bullet("Prompt 4 — GLOSS skin (Sonnet, High): ",
       "second UI descriptor, AMOUNT knob, A–D factory presets authored in-house.")

para("",
     "Order dependency: 1 → 2 → 3 → 4. Prompt 2 assumes the DSP class exists; Prompt 3 assumes "
     "the bridge is verified; Prompt 4 assumes the APEX UI and preset path exist.")

doc.save(OUT)
print("saved:", OUT)
